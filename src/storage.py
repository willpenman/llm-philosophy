"""Response storage helpers for append-only JSONL capture."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
import json
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

@dataclass(frozen=True)
class StoredText:
    path: Path
    text: str


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _slugify(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_").lower()
    return cleaned or "default"


def normalize_special_settings(value: str | None) -> str:
    if value is None or not str(value).strip():
        return "default"
    if str(value).strip().lower() == "default":
        return "default"
    return _slugify(str(value))


def format_input_text(system_text: str, puzzle_text: str) -> str:
    return f"System:\n{system_text}\n\nUser:\n{puzzle_text}"


def _format_filename_timestamp(created_at: str) -> str:
    timestamp = datetime.fromisoformat(created_at)
    timestamp = timestamp.astimezone(timezone.utc)
    return timestamp.strftime("%Y-%m-%dT%H%M%SZ")


def _format_display_date(created_at: str) -> str:
    timestamp = datetime.fromisoformat(created_at)
    timestamp = timestamp.astimezone(timezone.utc)
    return timestamp.strftime("%b %d, %Y")


def _docx_filename(
    special_settings: str,
    puzzle_name: str,
    puzzle_version: str | None,
    created_at: str,
) -> str:
    version = puzzle_version or "unknown"
    settings = normalize_special_settings(special_settings)
    timestamp = _format_filename_timestamp(created_at)
    return f"{settings}-{puzzle_name}-v{version}-{timestamp}.docx"


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.append(field)


def _add_text_paragraphs(document: Document, text: str) -> None:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    for line in normalized.split("\n"):
        # "Normal" is default, but this keeps us explicit
        document.add_paragraph(line, style="Normal")


def write_response_docx(
    *,
    path: Path,
    puzzle_prefix: str,
    display_name: str,
    puzzle_version: str | None,
    model_display: str,
    provider_display: str,
    settings_display: str,
    display_date: str,
    input_text: str,
    output_text: str,
) -> None:
    document = Document()
    document.core_properties.title = display_name
    # remove default 'after paragraph' spacing of 10pt, since models typically use 2 line breaks
    for style_name in ("Normal",):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(0)

    title_paragraph = document.add_paragraph(display_name, style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version_suffix = f" (v{puzzle_version})" if puzzle_version else ""
    headers = [
        f"{puzzle_prefix}: {display_name}{version_suffix}",
        f"LLM: {model_display} ({provider_display}){settings_display}",
        f"Completed: {display_date}",
    ]
    for header in headers:
        paragraph = document.add_paragraph(header)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)

    document.add_heading(f"Input given to {model_display}", level=1)
    _add_text_paragraphs(document, input_text)
    document.add_page_break()
    document.add_heading(f"{model_display}'s Output", level=1)
    _add_text_paragraphs(document, output_text)

    footer = document.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    _add_page_number(footer_paragraph)

    document.save(path)


def _base_record(
    *,
    run_id: str,
    created_at: str,
    provider: str,
    model: str,
    puzzle_name: str,
    puzzle_version: str | None,
    special_settings: str,
) -> dict[str, Any]:
    return {
        "run_id": run_id,
        "created_at": created_at,
        "provider": provider,
        "model": model,
        "puzzle_name": puzzle_name,
        "puzzle_version": puzzle_version,
        "special_settings": special_settings,
    }


class ResponseStore:
    """Append-only JSONL storage for requests, responses, and raw text."""

    def __init__(self, base_dir: Path) -> None:
        self.base_dir = base_dir

    def _provider_dir(self, provider: str, model: str) -> Path:
        return self.base_dir / provider / model

    def _append_jsonl(self, path: Path, record: dict[str, Any]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(record, ensure_ascii=True))
            handle.write("\n")

    def record_request(
        self,
        *,
        run_id: str,
        created_at: str,
        provider: str,
        model: str,
        puzzle_name: str,
        puzzle_version: str | None,
        special_settings: str,
        request_payload: dict[str, Any],
    ) -> Path:
        record = _base_record(
            run_id=run_id,
            created_at=created_at,
            provider=provider,
            model=model,
            puzzle_name=puzzle_name,
            puzzle_version=puzzle_version,
            special_settings=special_settings,
        )
        record["request"] = request_payload
        path = self._provider_dir(provider, model) / "requests.jsonl"
        self._append_jsonl(path, record)
        return path

    def record_response(
        self,
        *,
        run_id: str,
        created_at: str,
        provider: str,
        model: str,
        model_alias: str | None = None,
        provider_alias: str | None = None,
        puzzle_name: str,
        puzzle_title_prefix: str | None = None,
        puzzle_version: str | None,
        puzzle_title: str | None = None,
        special_settings: str,
        special_settings_display: str | None = None,
        request_payload: dict[str, Any],
        response_payload: dict[str, Any],
        input_text: str,
        output_text: str,
        derived: dict[str, Any] | None = None,
    ) -> StoredText:
        record = _base_record(
            run_id=run_id,
            created_at=created_at,
            provider=provider,
            model=model,
            puzzle_name=puzzle_name,
            puzzle_version=puzzle_version,
            special_settings=special_settings,
        )
        record["request"] = request_payload
        record["response"] = response_payload
        if derived:
            record["derived"] = derived
        response_path = self._provider_dir(provider, model) / "responses.jsonl"
        self._append_jsonl(response_path, record)

        text_dir = self._provider_dir(provider, model) / "texts"
        text_dir.mkdir(parents=True, exist_ok=True)
        filename = _docx_filename(
            special_settings=special_settings,
            puzzle_name=puzzle_name,
            puzzle_version=puzzle_version,
            created_at=created_at,
        )
        text_path = text_dir / filename
        display_name = puzzle_title or puzzle_name
        model_display = model_alias or model
        provider_display = provider_alias or provider
        settings_display = (
            ""
            if normalize_special_settings(special_settings) == "default"
            else f", {special_settings_display or special_settings}"
        )
        display_date = _format_display_date(created_at)
        puzzle_prefix = puzzle_title_prefix or "Philosophy problem"
        text_body = "\n".join(
            [
                f"{puzzle_prefix}: {display_name}",
                f"Model: {model_display} ({provider_display}){settings_display}",
                f"Completed: {display_date}",
                "",
                "---- INPUT ----",
                input_text,
                "",
                f"---- {model_display}'S OUTPUT ----",
                output_text,
            ]
        )
        write_response_docx(
            path=text_path,
            puzzle_prefix=puzzle_prefix,
            display_name=display_name,
            puzzle_version=puzzle_version,
            model_display=model_display,
            provider_display=provider_display,
            settings_display=settings_display,
            display_date=display_date,
            input_text=input_text,
            output_text=output_text,
        )
        return StoredText(path=text_path, text=text_body)
