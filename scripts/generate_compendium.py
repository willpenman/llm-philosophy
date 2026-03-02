"""Generate a combined Word document with all responses across all puzzles."""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.puzzles import load_puzzle
from src.response_reader import (
    display_names,
    extract_input_text,
    extract_output_text,
)
from src.storage import _format_display_date, normalize_special_settings

ROOT = Path(__file__).resolve().parents[1]


@dataclass
class ResponseRecord:
    """A parsed response record with extracted text."""

    provider: str
    model: str
    puzzle_name: str
    puzzle_version: str | None
    created_at: str
    special_settings: str
    system_text: str
    user_text: str
    output_text: str


def _parse_version(version: str | None) -> tuple[int, ...]:
    """Parse version string into comparable tuple."""
    if not version:
        return (0,)
    try:
        return tuple(int(x) for x in version.split("."))
    except ValueError:
        return (0,)


def load_all_responses(responses_dir: Path) -> list[ResponseRecord]:
    """Load all response records from JSONL files."""
    records: list[ResponseRecord] = []

    for response_path in responses_dir.glob("*/*/responses.jsonl"):
        with response_path.open(encoding="utf-8") as handle:
            for line in handle:
                line = line.strip()
                if not line:
                    continue
                data = json.loads(line)

                provider = data.get("provider")
                model = data.get("model")
                puzzle_name = data.get("puzzle_name")
                puzzle_version = data.get("puzzle_version")
                created_at = data.get("created_at")
                special_settings = normalize_special_settings(data.get("special_settings"))
                request_payload = data.get("request", {})
                response_payload = data.get("response", {})

                if not all([provider, model, puzzle_name, created_at]):
                    continue

                system_text, user_text = extract_input_text(provider, request_payload)
                output_text = extract_output_text(provider, response_payload)

                records.append(
                    ResponseRecord(
                        provider=provider,
                        model=model,
                        puzzle_name=puzzle_name,
                        puzzle_version=puzzle_version,
                        created_at=created_at,
                        special_settings=special_settings,
                        system_text=system_text,
                        user_text=user_text,
                        output_text=output_text,
                    )
                )

    return records


def select_best_responses(
    records: list[ResponseRecord],
) -> dict[str, dict[str, ResponseRecord]]:
    """Select the best response per (puzzle, model) pair.

    Prioritizes highest version, then most recent timestamp.
    Returns {puzzle_name: {model: record}}.
    """
    best: dict[str, dict[str, ResponseRecord]] = {}

    for record in records:
        puzzle = record.puzzle_name
        model = record.model

        if puzzle not in best:
            best[puzzle] = {}

        existing = best[puzzle].get(model)
        if existing is None:
            best[puzzle][model] = record
        else:
            # Compare by version first, then by timestamp
            existing_version = _parse_version(existing.puzzle_version)
            new_version = _parse_version(record.puzzle_version)

            if new_version > existing_version:
                best[puzzle][model] = record
            elif new_version == existing_version and record.created_at > existing.created_at:
                best[puzzle][model] = record

    return best


def _make_bookmark_name(text: str) -> str:
    """Create a valid bookmark name from text.

    Bookmark names must start with a letter and contain only letters, digits, and underscores.
    """
    import re
    # Replace non-alphanumeric with underscore, collapse multiple underscores
    name = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    # Ensure starts with letter
    if name and not name[0].isalpha():
        name = "bm_" + name
    return name or "bookmark"


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    """Add a bookmark around the entire paragraph content."""
    element = paragraph._p

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    element.insert(0, start)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    element.append(end)


def _add_internal_hyperlink(paragraph, text: str, bookmark_name: str) -> None:
    """Add an internal hyperlink to a bookmark within the document."""
    # Create hyperlink element with anchor (internal link)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    # Create run with text
    run = OxmlElement("w:r")

    # Add run properties for blue underlined link style
    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # Standard hyperlink blue
    rPr.append(color)

    # Underline
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)

    # Add text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_number(paragraph) -> None:
    """Add page number field to paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.append(field)


def _add_text_paragraphs(
    document: Document, text: str, *, page_break_after: bool = False
) -> None:
    """Add text paragraphs, treating 'System' and 'User' as h2 headings.

    Args:
        document: The Word document to add paragraphs to.
        text: The text content to add.
        page_break_after: If True, add a page break to the last paragraph instead
            of creating a separate page break paragraph. Avoids blank pages.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    # Remove trailing empty lines when adding page break to avoid blank pages
    if page_break_after:
        while lines and lines[-1] == "":
            lines.pop()
    last_para = None
    for line in lines:
        if line in ("System", "User"):
            last_para = document.add_heading(line, level=2)
        else:
            last_para = document.add_paragraph(line, style="Normal")
    if page_break_after and last_para is not None:
        last_para.add_run().add_break(WD_BREAK.PAGE)


def _set_section_header(section, text: str) -> None:
    """Set the running header for a section."""
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = header.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_compendium(responses_dir: Path, puzzles_dir: Path, output_path: Path) -> int:
    """Generate the combined compendium document.

    Returns the number of responses included.
    """
    records = load_all_responses(responses_dir)
    best_by_puzzle = select_best_responses(records)

    if not best_by_puzzle:
        print("No responses found.")
        return 0

    document = Document()
    document.core_properties.title = "LLM Philosophy Compendium"

    # Remove default paragraph spacing
    for style_name in ("Normal",):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.paragraph_format.space_after = Pt(0)

    # Title page
    title_para = document.add_paragraph("LLM Philosophy Compendium", style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("Complete Responses to Philosophy Problems")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    date_para = document.add_paragraph(f"Generated: {generated_date}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Build TOC structure first to know all entries
    toc_entries: list[tuple[str, str, int]] = []  # (display_text, bookmark_name, level)
    bookmark_id = 0

    sorted_puzzles = sorted(best_by_puzzle.keys())
    for puzzle_name in sorted_puzzles:
        models_dict = best_by_puzzle[puzzle_name]
        try:
            puzzle = load_puzzle(puzzle_name, puzzles_dir)
            puzzle_title = puzzle.title or puzzle_name
        except Exception:
            puzzle_title = puzzle_name

        puzzle_bookmark = _make_bookmark_name(f"puzzle_{puzzle_name}")
        toc_entries.append((puzzle_title, puzzle_bookmark, 1))

        sorted_models = sorted(
            models_dict.keys(),
            key=lambda m: display_names(models_dict[m].provider, m)[0],
        )
        for model in sorted_models:
            model_display, _ = display_names(models_dict[model].provider, model)
            model_bookmark = _make_bookmark_name(f"{puzzle_name}_{model}")
            toc_entries.append((model_display, model_bookmark, 2))

    # Table of Contents with manual hyperlinks
    document.add_heading("Table of Contents", level=1)
    for display_text, bookmark_name, level in toc_entries:
        toc_para = document.add_paragraph()
        if level == 2:
            # Indent level 2 entries
            toc_para.paragraph_format.left_indent = Pt(24)
        _add_internal_hyperlink(toc_para, display_text, bookmark_name)

    # Keep first section (title page + TOC) without running header
    # by not setting any header text on sections[0]

    # Now add content with bookmarks
    response_count = 0

    for puzzle_name in sorted_puzzles:
        models_dict = best_by_puzzle[puzzle_name]

        try:
            puzzle = load_puzzle(puzzle_name, puzzles_dir)
            puzzle_title = puzzle.title or puzzle_name
        except Exception:
            puzzle_title = puzzle_name

        puzzle_bookmark = _make_bookmark_name(f"puzzle_{puzzle_name}")

        # Add puzzle divider page (new section starts new page)
        new_section = document.add_section()
        _set_section_header(new_section, puzzle_title)

        puzzle_divider = document.add_paragraph(puzzle_title, style="Title")
        puzzle_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_bookmark(puzzle_divider, puzzle_bookmark, bookmark_id)
        bookmark_id += 1
        # No page break needed - next section will start a new page

        sorted_models = sorted(
            models_dict.keys(),
            key=lambda m: display_names(models_dict[m].provider, m)[0],
        )

        for model in sorted_models:
            record = models_dict[model]

            model_display, provider_display = display_names(record.provider, record.model)
            model_bookmark = _make_bookmark_name(f"{puzzle_name}_{model}")
            settings_display = (
                ""
                if record.special_settings == "default"
                else f", {record.special_settings}"
            )
            display_date = _format_display_date(record.created_at)
            version_suffix = f" (v{record.puzzle_version})" if record.puzzle_version else ""

            # Add new section for running header (one per response)
            header_text = f"{puzzle_title} — {model_display}"
            new_section = document.add_section()
            _set_section_header(new_section, header_text)

            # Model name as Title with bookmark
            title_para = document.add_paragraph(model_display, style="Title")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_bookmark(title_para, model_bookmark, bookmark_id)
            bookmark_id += 1

            # Metadata lines (centered, matching individual docx format)
            meta_lines = [
                f"Philosophy problem: {puzzle_title}{version_suffix}",
                f"LLM: {model_display} (by {provider_display}){settings_display}",
                f"Completed: {display_date}",
            ]
            for line in meta_lines:
                p = document.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)

            document.add_paragraph()  # spacing

            # Input section
            document.add_heading(f"Input given to {model_display}", level=1)
            input_text = f"System\n{record.system_text}\n\nUser\n{record.user_text}"
            _add_text_paragraphs(document, input_text, page_break_after=True)

            # Output section
            document.add_heading(f"{model_display}'s Output", level=1)
            _add_text_paragraphs(document, record.output_text)

            # No page break needed - next section will start a new page
            response_count += 1

    # Footer with page numbers
    for section in document.sections:
        footer = section.footer
        if footer.paragraphs:
            footer_para = footer.paragraphs[0]
            footer_para.text = ""
        else:
            footer_para = footer.add_paragraph()
        _add_page_number(footer_para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    return response_count


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a combined Word document with all responses."
    )
    parser.add_argument(
        "-o",
        "--output",
        default="responses/compendium.docx",
        help="Output file path (default: responses/compendium.docx)",
    )
    args = parser.parse_args()

    responses_dir = ROOT / "responses"
    puzzles_dir = ROOT / "prompts" / "puzzles"
    output_path = ROOT / args.output

    count = generate_compendium(responses_dir, puzzles_dir, output_path)
    print(f"Generated compendium with {count} responses: {output_path}")


if __name__ == "__main__":
    main()
